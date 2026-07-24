from __future__ import annotations

from io import BytesIO
from pathlib import Path


SUPPORTED_EXTENSIONS = {"txt", "md", "pdf", "docx"}


def decode_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore")


def extract_text_from_bytes(file_name: str, raw_bytes: bytes) -> str:
    extension = Path(file_name).suffix.lower().lstrip(".")
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: .{extension or 'unknown'}")
    if not raw_bytes:
        raise ValueError(f"{file_name} is empty")

    if extension in {"txt", "md"}:
        return decode_text(raw_bytes).strip()

    if extension == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw_bytes))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if not text:
            raise ValueError(f"Could not extract text from {file_name}. The PDF may be scanned.")
        return text

    from docx import Document

    document = Document(BytesIO(raw_bytes))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(f"Could not extract text from {file_name}")
    return text


def guess_person_name(cv_text: str, file_name: str) -> str:
    for line in cv_text.splitlines():
        cleaned = line.strip(" -•\t")
        if not cleaned:
            continue
        if len(cleaned) <= 80 and not any(char.isdigit() for char in cleaned):
            return cleaned
    return Path(file_name).stem.replace("_", " ").replace("-", " ").strip().title()
